"""
Document parser for RAG indexing and structured record analysis.

Supported formats: PDF, TXT, CSV, XLSX, HTML, DOCX.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any


class DocumentParser:
    """Parse documents into text chunks for RAG indexing."""

    def parse(self, file_path: str) -> list[str]:
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            chunks = self._parse_pdf(file_path)
        elif ext == ".txt":
            chunks = self._parse_txt(file_path)
        elif ext == ".csv":
            chunks = self._parse_csv(file_path)
        elif ext in (".xlsx", ".xls"):
            chunks = self._parse_xlsx(file_path)
        elif ext in (".html", ".htm"):
            chunks = self._parse_html(file_path)
        elif ext == ".docx":
            chunks = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext or 'unknown'}")

        chunks = [chunk for chunk in chunks if chunk and chunk.strip()]
        if not chunks:
            raise ValueError("No readable text found in this file.")
        return chunks

    def parse_records(self, file_path: str) -> list[dict[str, Any]]:
        ext = Path(file_path).suffix.lower()

        if ext == ".csv":
            records = self._csv_to_records(file_path)
        elif ext in (".xlsx", ".xls"):
            records = self._xlsx_to_records(file_path)
        else:
            raise ValueError(f"Records file must be CSV or XLSX, got: {ext or 'unknown'}")

        if not records:
            raise ValueError("No records found in this file.")
        return records

    # --- Text extraction ---

    def _parse_pdf(self, file_path: str) -> list[str]:
        import fitz

        doc = fitz.open(file_path)
        chunks: list[str] = []
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text().strip()
                if text:
                    for chunk in self._chunk_text(text):
                        chunks.append(f"[Page {page_num + 1}] {chunk}")
        finally:
            doc.close()
        return chunks

    def _parse_txt(self, file_path: str) -> list[str]:
        text = self._read_text_file(file_path)
        return self._chunk_text(text)

    def _parse_csv(self, file_path: str) -> list[str]:
        import pandas as pd

        df = pd.read_csv(file_path).fillna("")
        chunks = []
        for row_num, row in df.iterrows():
            row_text = " | ".join(
                f"{col}: {self._clean_cell(value)}" for col, value in row.items()
            )
            chunks.append(f"[Row {row_num + 1}] {row_text}")
        return chunks

    def _parse_xlsx(self, file_path: str) -> list[str]:
        import pandas as pd

        xls = pd.ExcelFile(file_path)
        chunks = []
        for sheet in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet).fillna("")
            for row_num, row in df.iterrows():
                row_text = " | ".join(
                    f"{col}: {self._clean_cell(value)}" for col, value in row.items()
                )
                chunks.append(f"[Sheet: {sheet}, Row {row_num + 1}] {row_text}")
        return chunks

    def _parse_html(self, file_path: str) -> list[str]:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(self._read_text_file(file_path), "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n").strip()
        return self._chunk_text(text)

    def _parse_docx(self, file_path: str) -> list[str]:
        from docx import Document

        doc = Document(file_path)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return self._chunk_text("\n".join(parts))

    # --- Structured records ---

    def _csv_to_records(self, file_path: str) -> list[dict[str, Any]]:
        import pandas as pd

        df = pd.read_csv(file_path).fillna("")
        return self._normalize_records(df.to_dict(orient="records"))

    def _xlsx_to_records(self, file_path: str) -> list[dict[str, Any]]:
        import pandas as pd

        df = pd.read_excel(file_path).fillna("")
        return self._normalize_records(df.to_dict(orient="records"))

    def _normalize_records(self, records: list[dict[str, Any]]) -> list[dict[str, Any]]:
        normalized = []
        for record in records:
            clean_record = {
                str(key).strip(): self._clean_cell(value)
                for key, value in record.items()
                if str(key).strip()
            }
            if any(str(value).strip() for value in clean_record.values()):
                normalized.append(clean_record)
        return normalized

    # --- Helpers ---

    def _read_text_file(self, file_path: str) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp1254", "latin-1"):
            try:
                with open(file_path, "r", encoding=encoding) as handle:
                    return handle.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file.")

    def _clean_cell(self, value: Any) -> str:
        if value is None:
            return ""
        return " ".join(str(value).strip().split())

    def _chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """Split text into overlapping word chunks."""
        if chunk_size <= overlap:
            raise ValueError("chunk_size must be greater than overlap.")

        words = text.split()
        if not words:
            return []

        chunks = []
        start = 0
        while start < len(words):
            end = min(start + chunk_size, len(words))
            chunks.append(" ".join(words[start:end]))
            if end == len(words):
                break
            start += chunk_size - overlap
        return chunks
